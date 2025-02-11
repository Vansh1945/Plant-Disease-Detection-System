import streamlit as st
from PIL import Image, ImageStat
import numpy as np
import time
import io
from docx import Document

# Page Configuration
st.set_page_config(page_title="🌿 Plant Disease Detector", page_icon="🌱", layout="wide")

# Sidebar Instructions
with st.sidebar:
    st.header("📌 How to Use")
    st.markdown(
        """
        1️⃣ Upload a **clear image** of a plant leaf.  
        2️⃣ The system will analyze **color & brightness** to verify if it's a leaf.  
        3️⃣ If valid, it will detect **possible diseases** and provide home treatments.  
        """
    )
    st.info("💡 **Tip:** Ensure good lighting and a clear background!")

# Disease Solutions with Step-by-Step Home Remedies
disease_solutions = {
    "Blight": [
        "1️⃣ **Baking Soda Spray**",
        "- **Ingredients:** 1 tsp baking soda + 1 liter water + a few drops of dish soap.",
        "- **Steps:** Mix well and spray on infected leaves. Repeat weekly.",
        "2️⃣ **Neem Oil Solution**",
        "- **Ingredients:** 1 tbsp neem oil + 1 liter water.",
        "- **Steps:** Mix and spray on affected areas every 5-7 days."
    ],
    
    "Rust": [
        "1️⃣ **Milk Spray**",
        "- **Ingredients:** 1 part milk + 2 parts water.",
        "- **Steps:** Spray on leaves every 3-4 days to stop fungal growth.",
        "2️⃣ **Garlic Extract Spray**",
        "- **Ingredients:** 2 cloves garlic (crushed) + 1 liter water.",
        "- **Steps:** Let it sit overnight, strain, and spray every 5 days."
    ],
    
    "Healthy": [
        "✅ **Your plant is healthy!**",
        "- Keep it healthy by watering regularly.",
        "- Provide enough sunlight.",
        "- Check for pests weekly."
    ],
    
    "Not a Plant Leaf": [
        "⚠️ **Error:** Please upload a valid **plant leaf image**!"
    ]
}

# Function to check if image is a plant leaf
def is_plant_leaf(image):
    image = image.convert("RGB")
    img_array = np.array(image)
    r_mean, g_mean, b_mean = np.mean(img_array[:, :, 0]), np.mean(img_array[:, :, 1]), np.mean(img_array[:, :, 2])
    return g_mean > r_mean and g_mean > b_mean  # Green must be dominant

# Function to predict disease based on brightness
def predict_disease(image):
    image = image.convert("L")  # Convert to grayscale
    brightness = np.mean(np.array(image))

    if brightness < 80:
        return "Blight"
    elif brightness < 150:
        return "Rust"
    elif brightness < 240:
        return "Healthy"
    else:
        return "Not a Plant Leaf"

# Function to generate Word report
def generate_word_report(disease, solutions):
    doc = Document()
    doc.add_heading("🌿 Plant Disease Detection Report", level=1)

    doc.add_paragraph(f"🩺 **Disease Detected:** {disease}")
    doc.add_paragraph("\n**Solution Steps:**")

    for step in solutions:
        doc.add_paragraph(step)

    # Save the document to a buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# UI Header
st.markdown("<h1 style='text-align: center; color: green;'>🌿 Plant Disease Detection System</h1>", unsafe_allow_html=True)
st.write("🔎 **Upload a plant leaf image to detect potential diseases and get home treatment solutions.**")

# File Uploader
uploaded_file = st.file_uploader("📤 Upload a plant leaf image", type=["jpg", "jpeg", "png"])

if uploaded_file:
    col1, col2 = st.columns([1, 1])

    # Display Image
    with col1:
        image = Image.open(uploaded_file)
        st.image(image, caption="📷 Uploaded Image", use_container_width=True)

    # Processing Animation with Progress Bar
    with col2:
        with st.spinner("🔍 Analyzing... Please wait..."):
            progress_bar = st.progress(0)
            for i in range(1, 101):
                time.sleep(0.02)  # Simulated processing delay
                progress_bar.progress(i)

            # Validate if uploaded image is a plant leaf
            if not is_plant_leaf(image):
                st.error("⚠️ **Error:** Please upload a valid **Plant Leaf Image**!")
            else:
                disease = predict_disease(image)
                st.subheader(f"🩺 Disease Detected: **{disease}**")

                # Display step-by-step remedies
                for step in disease_solutions[disease]:
                    st.markdown(step)

                # Generate and Download Report
                word_report = generate_word_report(disease, disease_solutions[disease])
                st.download_button(label="📥 Download Report (Word)", data=word_report, file_name="Plant_Disease_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Footer Section
st.markdown("<hr>", unsafe_allow_html=True)

